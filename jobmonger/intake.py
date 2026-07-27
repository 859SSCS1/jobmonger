"""[DOC-INTAKE] — read a document into memory. It is never copied to disk.

Supported: .txt, .md (standard library), .pdf (pypdf), .docx (python-docx).
Scanned-image OCR is deferred per the Stage 1 scope.

The two optional dependencies are imported lazily and, when missing, produce a
plain sentence telling the user what to install — not an ImportError traceback.
The person using this tool may be having a bad week; a stack trace is a poor way
to tell them their file type needs one more package.
"""

from __future__ import annotations

import unicodedata
from dataclasses import dataclass
from pathlib import Path

from . import consent, log

SUPPORTED_SUFFIXES = (".txt", ".md", ".markdown", ".pdf", ".docx")

# Generous, but bounded. A handbook is large; a 200 MB file is a mistake, and
# discovering that after a five-minute parse is a worse experience than being
# told immediately.
MAX_BYTES = 25 * 1024 * 1024


class IntakeError(Exception):
    """A document could not be read. The message is shown to the user as-is."""


@dataclass(frozen=True)
class Document:
    """A document held in memory. Never written to this tool's own storage."""

    source_name: str
    text: str
    #: 1-based page or paragraph-block boundaries, for showing "where" later.
    #: Each entry is (label, start_offset, end_offset) into ``text``.
    sections: tuple[tuple[str, int, int], ...] = ()

    @property
    def char_count(self) -> int:
        return len(self.text)

    @property
    def word_count(self) -> int:
        return len(self.text.split())


def load(path: str | Path) -> Document:
    """Read a document from disk into memory.

    Requires consent: reading a document is the first act of using the tool,
    and the gate belongs in front of it rather than in front of the model call.
    """
    consent.require()

    source = Path(path)
    if not source.exists():
        raise IntakeError(f"No file at {source}.")
    if not source.is_file():
        raise IntakeError(f"{source} is a folder, not a file.")

    size = source.stat().st_size
    if size == 0:
        raise IntakeError(f"{source.name} is empty.")
    if size > MAX_BYTES:
        mb = size / (1024 * 1024)
        raise IntakeError(
            f"{source.name} is {mb:.0f} MB, larger than the {MAX_BYTES // (1024 * 1024)} MB limit."
        )

    suffix = source.suffix.lower()
    if suffix in (".txt", ".md", ".markdown"):
        document = _read_text(source)
    elif suffix == ".pdf":
        document = _read_pdf(source)
    elif suffix == ".docx":
        document = _read_docx(source)
    else:
        supported = ", ".join(SUPPORTED_SUFFIXES)
        raise IntakeError(
            f"{source.name} is a {suffix or 'file with no extension'}. "
            f"Supported types are: {supported}."
        )

    if not document.text.strip():
        raise IntakeError(
            f"No text could be read from {source.name}. "
            "If it is a scan or a photo of a page, this version cannot read it yet."
        )

    # Filename only — never the path (which can contain the user's real name)
    # and never any content.
    log.record(
        "document.loaded",
        source_name=document.source_name,
        suffix=suffix,
        characters=document.char_count,
        words=document.word_count,
        sections=len(document.sections),
    )
    return document


def _normalise(text: str) -> str:
    """Make text safe to match against and pleasant to display.

    PDFs in particular arrive full of non-breaking spaces, soft hyphens, and
    ligatures. Left alone, these break name matching in ways that are invisible
    on screen — the redaction gate would miss a name the user can plainly see,
    which is the worst possible failure for this tool.
    """
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("­", "")  # soft hyphen
    text = text.replace(" ", " ")  # non-breaking space
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse runs of blank lines, but keep paragraph structure.
    lines = [line.rstrip() for line in text.split("\n")]
    out: list[str] = []
    blanks = 0
    for line in lines:
        if line:
            blanks = 0
            out.append(line)
        else:
            blanks += 1
            if blanks <= 2:
                out.append("")
    return "\n".join(out).strip()


def _assemble(parts: list[tuple[str, str]], source_name: str) -> Document:
    """Join labelled parts into one document, recording section offsets."""
    chunks: list[str] = []
    sections: list[tuple[str, int, int]] = []
    cursor = 0
    for label, body in parts:
        body = _normalise(body)
        if not body:
            continue
        if chunks:
            chunks.append("\n\n")
            cursor += 2
        start = cursor
        chunks.append(body)
        cursor += len(body)
        sections.append((label, start, cursor))
    return Document(source_name=source_name, text="".join(chunks), sections=tuple(sections))


def _read_text(source: Path) -> Document:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            raw = source.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
        except OSError as exc:
            raise IntakeError(f"Could not open {source.name}: {exc}") from exc
        return Document(source_name=source.name, text=_normalise(raw))
    raise IntakeError(f"Could not work out the text encoding of {source.name}.")


def _read_pdf(source: Path) -> Document:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise IntakeError(
            "Reading PDFs needs one extra package. Install it with:\n"
            "    pip install pypdf\n"
            "Or install everything at once with:  pip install jobmonger[docs]"
        ) from None

    try:
        reader = PdfReader(str(source))
    except Exception as exc:
        raise IntakeError(f"{source.name} could not be opened as a PDF: {exc}") from exc

    if getattr(reader, "is_encrypted", False):
        try:
            reader.decrypt("")
        except Exception:
            pass
        if getattr(reader, "is_encrypted", False):
            raise IntakeError(
                f"{source.name} is password-protected. Open it, save an unlocked copy, "
                "and try that instead."
            )

    parts: list[tuple[str, str]] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            parts.append((f"Page {index}", page.extract_text() or ""))
        except Exception:
            # One unreadable page should not lose the other ninety.
            parts.append((f"Page {index}", ""))
    return _assemble(parts, source.name)


def _read_docx(source: Path) -> Document:
    try:
        import docx  # python-docx
    except ImportError:
        raise IntakeError(
            "Reading Word documents needs one extra package. Install it with:\n"
            "    pip install python-docx\n"
            "Or install everything at once with:  pip install jobmonger[docs]"
        ) from None

    try:
        document = docx.Document(str(source))
    except Exception as exc:
        raise IntakeError(f"{source.name} could not be opened as a Word document: {exc}") from exc

    body = [p.text for p in document.paragraphs]

    # Tables carry a lot of what matters in HR documents — pay bands, leave
    # accruals, escalation steps. Dropping them would quietly lose the answer
    # to the question the user most likely came here to ask.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                body.append(" | ".join(cells))

    return _assemble([("Document", "\n".join(body))], source.name)


def from_text(text: str, source_name: str = "pasted text") -> Document:
    """Wrap text the user pasted directly, bypassing the filesystem."""
    consent.require()
    cleaned = _normalise(text)
    if not cleaned.strip():
        raise IntakeError("Nothing to read — the text was empty.")
    log.record(
        "document.loaded",
        source_name=source_name,
        suffix="(pasted)",
        characters=len(cleaned),
        words=len(cleaned.split()),
        sections=0,
    )
    return Document(source_name=source_name, text=cleaned)
